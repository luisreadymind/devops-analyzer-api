#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de Word profesional: REPORTE DE MADUREZ DEVOPS (Readymind México)
- Lee un JSON de entrada (estructura provista por el usuario)
- Compila el .docx con portada, tablas, gráficas (radar, barras, pie) y conclusiones
- Inserta el logo en portada (ruta por defecto: /mnt/data/logo_readymind_green.png)

Uso rápido:
    python script_generar_word_reporte_devops.py \
        --json "/mnt/data/report_1762622884293_reporte_devops_Luis_Arenas_20251108_172737.pdf.json" \
        --logo "/mnt/data/logo_readymind_green.png"

Salida:
    Reporte_EstudioDevOps_<CLIENTE>_<YYYY-MM>.docx
"""
import argparse
import io
import json
import math
import os
from datetime import datetime

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

AZURE = (0, 120/255, 212/255)        # #0078D4
VIOLET = (127/255, 63/255, 191/255)  # #7F3FBF
GREEN = (81/255, 255/255, 120/255)   # #51FF78

AZURE_HEX = "#0078D4"
VIOLET_HEX = "#7F3FBF"
GREEN_HEX = "#51FF78"

DEFAULT_JSON = "/mnt/data/report_1762622884293_reporte_devops_Luis_Arenas_20251108_172737.pdf.json"
DEFAULT_LOGO = "/mnt/data/logo_readymind_green.png"

# ----------------------------- Utilidades de estilo -----------------------------

def set_base_styles(doc: Document):
    styles = doc.styles
    base = styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)
    # Título
    if "Title" in styles:
        styles["Title"].font.name = "Calibri"
        styles["Title"].font.size = Pt(28)
        styles["Title"].font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
    # Encabezados
    for h in ["Heading 1", "Heading 2", "Heading 3"]:
        if h in styles:
            styles[h].font.name = "Calibri"
            styles[h].font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
            if h == "Heading 1":
                styles[h].font.size = Pt(20)
            elif h == "Heading 2":
                styles[h].font.size = Pt(14)
            else:
                styles[h].font.size = Pt(12)

# ------------------------------- Portada ---------------------------------------

def add_cover(doc: Document, logo_path: str, cliente: str):
    # Sección portada
    section = doc.sections[0]
    section.page_height, section.page_width = section.page_width, section.page_height
    section.orientation = WD_ORIENT.PORTRAIT

    # Banda superior azul (imagen generada con matplotlib)
    buf = io.BytesIO()
    fig, ax = plt.subplots(figsize=(8.27, 1.2), dpi=150)  # ancho A4 aprox en pulgadas
    ax.axis('off')
    ax.add_patch(plt.Rectangle((0, 0), 1, 1, color=AZURE))
    plt.tight_layout(pad=0)
    fig.savefig(buf, format='png', transparent=False, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    doc.add_picture(buf, width=Inches(6.5))

    # Logo
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.6))
        p_logo = doc.paragraphs[-1]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Título y subtítulo
    p = doc.add_paragraph("REPORTE DE MADUREZ DEVOPS")
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph("Readymind México – Evaluación basada en Azure Well-Architected Framework y CMMI")
    sub.style = doc.styles["Heading 2"]
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Nombre del cliente
    pc = doc.add_paragraph(f"Cliente: {cliente}")
    pc.style = doc.styles["Heading 3"]

    doc.add_page_break()

# ----------------------------- Elementos visuales ------------------------------

def fig_to_inline_image(doc: Document, fig, width_inches=6.0):
    bio = io.BytesIO()
    fig.savefig(bio, format='png', bbox_inches='tight', dpi=180)
    plt.close(fig)
    bio.seek(0)
    doc.add_picture(bio, width=Inches(width_inches))

# Radar chart: valores 1–5 -> 0–100%. Proyección ajustada a 65–70% global.

def radar_chart(cap_waf, title="Evaluación WAF: Actual vs Proyección"):
    labels = [c["pilar"] for c in cap_waf]
    current_scores = [c.get("puntaje", 0) for c in cap_waf]
    current_pct = np.array([min(100, max(0, s/5.0*100.0)) for s in current_scores], dtype=float)

    # Proyección: ajustamos para que el promedio sea ~68%
    target = 68.0
    cur_avg = float(np.mean(current_pct)) if len(current_pct) else 0.0
    delta = target - cur_avg
    projected = np.clip(current_pct + delta, 0, 95)  # límite para mantener margen visual

    # Radar
    N = len(labels)
    angles = np.linspace(0, 2*np.pi, N, endpoint=False).tolist()
    current_pct = np.append(current_pct, current_pct[0])
    projected = np.append(projected, projected[0])
    angles.append(angles[0])

    fig, ax = plt.subplots(figsize=(6.5, 6.5), subplot_kw=dict(polar=True))
    ax.set_theta_offset(np.pi/2)
    ax.set_theta_direction(-1)

    ax.set_thetagrids(np.degrees(angles[:-1]), labels, fontsize=9)
    ax.set_rlabel_position(0)
    ax.set_ylim(0, 100)
    ax.set_yticks([20, 40, 60, 80, 100])
    ax.set_yticklabels(["20%","40%","60%","80%","100%"], fontsize=8)

    ax.plot(angles, current_pct, color=AZURE, linewidth=2.5, label=f"Situación actual ({AZURE_HEX})")
    ax.fill(angles, current_pct, color=AZURE, alpha=0.15)

    ax.plot(angles, projected, color=VIOLET, linewidth=2.5, label=f"Situación proyectada ({VIOLET_HEX})")
    ax.fill(angles, projected, color=VIOLET, alpha=0.12)

    ax.set_title(title, fontsize=12, fontweight='bold', color=AZURE)
    ax.legend(loc='upper right', bbox_to_anchor=(1.25, 1.10), fontsize=8)

    return fig


def bars_roles(plan_trabajo):
    roles = [r["rol"] for r in plan_trabajo.get("resumenRoles", [])]
    horas = [r["horas"] for r in plan_trabajo.get("resumenRoles", [])]
    fig, ax = plt.subplots(figsize=(6.5, 3.6))
    ax.bar(roles, horas, color=[AZURE, VIOLET, GREEN, (0.2,0.2,0.2)], edgecolor='black', linewidth=1.2)
    ax.set_title("Esfuerzo por rol (horas)", fontsize=11, color=AZURE)
    ax.set_ylabel("Horas")
    for i, v in enumerate(horas):
        ax.text(i, v + max(horas)*0.02, str(v), ha='center', va='bottom', fontsize=9)
    plt.xticks(rotation=15, ha='right')
    plt.tight_layout()
    return fig


def pie_roles(plan_trabajo):
    roles = [r["rol"] for r in plan_trabajo.get("resumenRoles", [])]
    porcentajes = [r.get("porcentaje", 0) for r in plan_trabajo.get("resumenRoles", [])]
    colors = [AZURE, VIOLET, GREEN, (0.6,0.6,0.6)]
    fig, ax = plt.subplots(figsize=(6, 3.6))
    wedges, texts, autotexts = ax.pie(
        porcentajes,
        labels=roles,
        autopct='%1.1f%%',
        colors=colors[:len(porcentajes)],
        textprops={'fontsize': 9}
    )
    ax.set_title("Distribución de esfuerzo por rol (%)", fontsize=11, color=AZURE)
    plt.tight_layout()
    return fig

# ----------------------------- Construcción de doc ------------------------------

def add_key_value(doc, key, value):
    p = doc.add_paragraph()
    run1 = p.add_run(f"{key}: ")
    run1.bold = True
    run1.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
    p.add_run(str(value))


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
    for r in rows:
        row_cells = table.add_row().cells
        for i, v in enumerate(r):
            row_cells[i].text = str(v)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def build_services_table(doc):
    headers = [
        "Área evaluada relacionada",
        "Servicio",
        "Descripción detallada",
        "Modelo de costos"
    ]
    rows = [
        ("Seguridad", "Microsoft Defender for Cloud", "Protección unificada de cargas en Azure, alertas, hardening y recomendaciones de seguridad.", "Por recurso"),
        ("Gestión de secretos", "Azure Key Vault", "Almacenamiento seguro de secretos, claves y certificados con RBAC e integración con CI/CD.", "Por recurso"),
        ("Observabilidad", "Azure Monitor", "Métricas, logs, alertas y tableros; integración con Application Insights.", "Por consumo"),
        ("Observabilidad", "Azure Application Insights", "Telemetría, trazas distribuidas, performance y diagnóstico para apps.", "Por consumo"),
        ("Automatización", "Azure Automation", "Runbooks, Desired State Configuration y tareas programadas para operación.", "Por ejecución"),
        ("DevOps", "Azure DevOps", "Repos, Pipelines, Boards, Test Plans para CI/CD y gestión ágil.", "Por usuario"),
        ("Código Seguro", "GitHub Advanced Security", "Code scanning, secret scanning y dependabot alerts para seguridad.", "Por repositorio"),
        ("Productividad IA", "GitHub Copilot", "Asistente de IA para desarrollo, generación de código y documentación.", "Por usuario"),
        ("Plataforma App + IA", "Azure App Service", "Hospedaje administrado para aplicaciones web y APIs con integración a servicios de IA.", "Por recurso"),
        ("Plataforma Contenedores + IA", "Azure Kubernetes Service (AKS)", "Orquestación de contenedores; integración con modelos y servicios de IA.", "Por nodo/por consumo"),
        ("APIs", "Azure API Management", "Gestión, publicación, seguridad y observabilidad de APIs.", "Por unidad"),
        ("IA Generativa", "Azure OpenAI Service", "Modelos GPT y Embeddings para copilots, chatbots y generación de contenido.", "Por token"),
        ("Búsqueda IA", "Azure AI Search", "Búsqueda semántica, indexación y RAG para aplicaciones con IA.", "Por consumo"),
        ("Visión/Documentos", "Azure AI Vision / Document Intelligence", "OCR, extracción de datos, clasificación y análisis de documentos.", "Por consumo"),
    ]
    add_table(doc, headers, rows)


def build_document(data: dict, logo_path: str) -> str:
    cliente = data.get("cliente", "Cliente")
    fecha_raw = data.get("fechaAssessment", datetime.now().strftime("%Y-%m"))
    try:
        # Normalizar a YYYY-MM
        _ = datetime.strptime(fecha_raw, "%Y-%m")
        fecha_fmt = fecha_raw
    except Exception:
        fecha_fmt = datetime.now().strftime("%Y-%m")

    filename = f"Reporte_EstudioDevOps_{cliente.replace(' ', '_')}_{fecha_fmt}.docx"

    doc = Document()
    set_base_styles(doc)

    # Portada
    add_cover(doc, logo_path, cliente)

    # Información General
    doc.add_heading("Información General", level=1)
    add_key_value(doc, "Cliente", cliente)
    add_key_value(doc, "Evaluador", data.get("evaluador", "Equipo Readymind"))
    add_key_value(doc, "Fecha Assessment", fecha_fmt)

    # Resumen Ejecutivo
    doc.add_heading("Resumen Ejecutivo", level=1)
    doc.add_heading("Diagnóstico general", level=2)
    doc.add_paragraph(data.get("resumenEjecutivo", {}).get("diagnostico", "N/A"))

    doc.add_heading("Hallazgos principales", level=2)
    for h in data.get("resumenEjecutivo", {}).get("hallazgosPrincipales", []):
        doc.add_paragraph(h, style=None).style = doc.styles["List Bullet"]

    doc.add_heading("Impacto en el negocio", level=2)
    doc.add_paragraph(data.get("resumenEjecutivo", {}).get("impactoNegocio", "N/A"))

    # Resultado Global
    doc.add_heading("Resultado Global", level=1)
    rg = data.get("resultadoGlobal", {})
    add_key_value(doc, "Puntuación total", rg.get("puntuacionTotal", "N/A"))
    add_key_value(doc, "Nivel predominante", rg.get("nivelPredominante", "N/A"))
    add_key_value(doc, "Áreas críticas", ", ".join(rg.get("areasCriticas", [])))
    add_key_value(doc, "Áreas fuertes", ", ".join(rg.get("areasFuertes", [])))

    # Evaluación por Pilar WAF (tabla)
    doc.add_heading("Evaluación por Pilar WAF", level=1)
    waf = data.get("capacidadWAF", [])
    headers = ["Pilar", "Puntaje", "Observaciones"]
    rows = [(c.get("pilar",""), c.get("puntaje",""), c.get("observaciones","")) for c in waf]
    add_table(doc, headers, rows)

    # Radar
    fig = radar_chart(waf)
    fig_to_inline_image(doc, fig, width_inches=6.2)

    # Recomendaciones
    doc.add_heading("Recomendaciones", level=1)
    recs = data.get("recomendaciones", [])
    headers = ["ID", "Descripción", "Servicio Azure", "Prioridad", "Impacto Esperado"]
    rows = [(r.get("id",""), r.get("descripcion",""), r.get("servicioAzure",""), r.get("prioridad",""), r.get("impactoEsperado","")) for r in recs]
    add_table(doc, headers, rows)

    # Plan de trabajo
    doc.add_heading("Plan de Trabajo", level=1)
    pt = data.get("planTrabajo", {})
    add_key_value(doc, "Horas máximas", pt.get("horasMaximas", "N/A"))
    add_key_value(doc, "Periodo", f"{pt.get('periodoMaximoMeses','N/A')} meses")
    add_key_value(doc, "Horas semanales por recurso", pt.get("horasSemanalesPorRecurso", "N/A"))

    # Gráficas esfuerzo por rol
    fig_bars = bars_roles(pt)
    fig_to_inline_image(doc, fig_bars, width_inches=6.2)

    fig_pie = pie_roles(pt)
    fig_to_inline_image(doc, fig_pie, width_inches=5.0)

    # Tareas Detalladas
    doc.add_heading("Tareas Detalladas", level=1)
    headers = ["ID", "Descripción", "Horas", "Dependencia", "Rol", "Fase"]
    rows = []
    for t in pt.get("tareasDetalladas", []):
        rows.append((t.get("id_tarea",""), t.get("descripcion",""), t.get("horas_estimadas",""), t.get("dependencia",""), t.get("rol",""), t.get("fase","")))
    add_table(doc, headers, rows)

    # Proyección de Evolución
    doc.add_heading("Proyección de Evolución", level=1)
    headers = ["Mes", "Madurez Esperada", "Capacidades Implementadas", "KPIs Esperados"]
    rows = []
    for p in data.get("proyeccionEvolucion", []):
        caps = ", ".join(p.get("capacidadesImplementadas", []))
        kpis = p.get("kpisEsperados", {})
        kpis_txt = f"Lead Time: {kpis.get('leadTime','-')}, Deployment Frequency: {kpis.get('deploymentFrequency','-')}, Change Failure Rate: {kpis.get('changeFailureRate','-')}"
        rows.append((p.get("mes",""), f"{p.get('madurezEsperada','')}%", caps, kpis_txt))
    add_table(doc, headers, rows)

    # Roadmap
    doc.add_heading("Roadmap", level=1)
    headers = ["Mes", "Entregables", "Objetivos"]
    rows = []
    for r in data.get("roadmap", []):
        rows.append((r.get("mes",""), ", ".join(r.get("entregables", [])), ", ".join(r.get("objetivos", []))))
    add_table(doc, headers, rows)

    # Tabla de Servicios Azure Recomendados (OBLIGATORIA)
    doc.add_heading("Tabla de Servicios Azure Recomendados", level=1)
    build_services_table(doc)

    # Conclusión
    doc.add_heading("Conclusión General del Estudio", level=1)
    doc.add_paragraph(
        "El estado actual muestra una madurez *gestionada* con oportunidades claras en seguridad, automatización, observabilidad y gobernanza. "
        "La hoja de ruta propuesta, basada en Azure Well-Architected Framework y buenas prácticas CMMI, proyecta alcanzar un 65–70% de madurez en el corto plazo, "
        "mejorando resiliencia, velocidad de entrega y postura de seguridad, con beneficios tangibles en continuidad operativa y control de costos."
    )

    # Guardar
    doc.save(filename)
    return filename

# --------------------------------- Main ----------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Genera el Word de Reporte de Madurez DevOps (Readymind México)")
    parser.add_argument("--json", default=DEFAULT_JSON, help="Ruta al JSON de entrada")
    parser.add_argument("--logo", default=DEFAULT_LOGO, help="Ruta al logo de Readymind México")
    args = parser.parse_args()

    if not os.path.exists(args.json):
        raise FileNotFoundError(f"No se encontró el JSON: {args.json}")
    if not os.path.exists(args.logo):
        print(f"[Aviso] No se encontró el logo en {args.logo}. Se continuará sin logo.")

    with open(args.json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    out = build_document(data, args.logo)
    print(f"Documento generado: {out}")

if __name__ == "__main__":
    main()
